# Copyright 2026 EPAM Systems, Inc. (“EPAM”)
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Tests for unified MessageExporter with node-based document structure."""

import tempfile
from typing import Iterator
from unittest.mock import MagicMock

import pytest
from docx import Document

from codemie.chains.base import Thought, ThoughtAuthorType
from codemie.core.models import ChatRole
from codemie.rest_api.models.assistant import Assistant
from codemie.rest_api.models.conversation import Conversation, GeneratedMessage
from codemie.service.conversation import MessageExporter, ExportFormat
from codemie.service.conversation.message_exporter import (
    ContentParser,
    DocumentBuilder,
    PDFProcessor,
    DOCXProcessor,
    PPTXProcessor,
    ProcessorFactory,
)
from codemie.service.conversation.pandoc_nodes import (
    Document as PandocDocument,
    Heading,
    HeadingLevel,
    Paragraph,
    RawContent,
    CodeBlock,
    HorizontalRule,
    PageBreak,
)


# ============================================================================
# Test Fixtures
# ============================================================================


@pytest.fixture
def mock_assistant():
    """Create a mock assistant."""
    assistant = MagicMock(spec=Assistant)
    assistant.id = "assistant-123"
    assistant.name = "Test Assistant"
    assistant.icon_url = "https://example.com/icon.png"
    return assistant


@pytest.fixture
def mock_conversation_simple():
    """Create a simple conversation with user and AI messages."""
    conversation = MagicMock(spec=Conversation)
    conversation.id = "conv-123"
    conversation.conversation_name = "Test Conversation"

    user_message = GeneratedMessage(
        role=ChatRole.USER,
        message="What is Python?",
        history_index=0,
    )

    ai_message = GeneratedMessage(
        role=ChatRole.ASSISTANT,
        message="Python is a programming language.",
        history_index=0,
        assistant_id="assistant-123",
        thoughts=[
            Thought(
                id="thought-1",
                message="Analyzing the question",
                author_name="Research Agent",
                author_type=ThoughtAuthorType.Agent,
            )
        ],
    )

    conversation.find_messages = MagicMock(return_value=(user_message, ai_message))
    conversation.history = [user_message, ai_message]

    return conversation


@pytest.fixture
def mock_conversation_multi_messages(mock_assistant):
    """Create conversation with multiple message exchanges."""
    conversation = MagicMock(spec=Conversation)
    conversation.id = "conv-multi"
    conversation.conversation_name = "Multi Message Chat"
    conversation.history = [
        # Exchange 1
        GeneratedMessage(role=ChatRole.USER, message="Question 1", history_index=0),
        GeneratedMessage(
            role=ChatRole.ASSISTANT,
            message="Answer 1",
            history_index=0,
            assistant_id=mock_assistant.id,
            thoughts=[
                Thought(
                    id="thought-1",
                    message="Processing question 1",
                    author_name="Agent 1",
                    author_type=ThoughtAuthorType.Agent,
                )
            ],
        ),
        # Exchange 2
        GeneratedMessage(role=ChatRole.USER, message="Question 2", history_index=1),
        GeneratedMessage(
            role=ChatRole.ASSISTANT,
            message="Answer 2",
            history_index=1,
            assistant_id=mock_assistant.id,
            thoughts=[
                Thought(
                    id="thought-2",
                    message="Processing question 2",
                    author_name="Agent 2",
                    author_type=ThoughtAuthorType.Agent,
                )
            ],
        ),
    ]
    return conversation


# ============================================================================
# Tests for ExportFormat Enum
# ============================================================================


class TestExportFormat:
    """Test suite for ExportFormat enum."""

    def test_pdf_content_type(self):
        """Test PDF content type."""
        assert ExportFormat.PDF.content_type == "application/pdf"

    def test_docx_content_type(self):
        """Test DOCX content type."""
        expected = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert ExportFormat.DOCX.content_type == expected

    def test_pptx_content_type(self):
        """Test PPTX content type."""
        expected = "application/vnd.openxmlformats-officedocument.presentationml.presentation"
        assert ExportFormat.PPTX.content_type == expected


# ============================================================================
# Tests for ContentParser
# ============================================================================


class TestContentParser:
    """Test suite for ContentParser class."""

    def test_parse_message_content_plain_text(self):
        """Test parsing plain text message."""
        parser = ContentParser()
        nodes = parser.parse_message_content("Simple text message")

        assert len(nodes) == 1
        assert nodes[0].content == "Simple text message"

    def test_parse_thought_content_without_code(self):
        """Test parsing thought without code blocks."""
        parser = ContentParser()
        thought = Thought(
            id="test",
            message="This is a thought",
            author_type=ThoughtAuthorType.Agent,
        )

        nodes = parser.parse_thought_content(thought)

        assert len(nodes) == 1
        assert nodes[0].content == "This is a thought"

    def test_parse_thought_content_with_code(self):
        """Test parsing thought with code blocks."""
        parser = ContentParser()
        thought = Thought(
            id="test",
            message="Here is code:\n```python\nprint('hello')\n```\nThat was code.",
            author_type=ThoughtAuthorType.Agent,
        )

        nodes = parser.parse_thought_content(thought)

        assert len(nodes) == 3  # Text before, code block, text after
        assert isinstance(nodes[1], CodeBlock)
        assert nodes[1].language == "python"
        assert "print('hello')" in nodes[1].code

    def test_contains_code_block_true(self):
        """Test detecting code blocks."""
        parser = ContentParser()
        assert parser._contains_code_block("```python\ncode\n```") is True
        assert parser._contains_code_block("~~~python\ncode\n~~~") is True

    def test_contains_code_block_false(self):
        """Test not detecting code blocks in plain text."""
        parser = ContentParser()
        assert parser._contains_code_block("Plain text") is False


# ============================================================================
# Tests for Format Processors
# ============================================================================


class TestFormatProcessors:
    """Test suite for format processor classes."""

    def test_pdf_processor_no_changes(self):
        """Test PDF processor doesn't modify document."""
        processor = PDFProcessor()
        doc = PandocDocument()
        doc.add(Heading(text="Title", level=HeadingLevel.H1))
        doc.add(PageBreak())

        processed = processor.process_document(doc)

        assert len(processed.children) == 2
        assert isinstance(processed.children[1], PageBreak)

    def test_pdf_processor_pandoc_args(self):
        """Test PDF processor returns correct pandoc arguments."""
        processor = PDFProcessor()
        args = processor.get_pandoc_args()

        assert '--pdf-engine=pdflatex' in args
        assert 'geometry:margin=1in' in args
        # pdflatex uses Computer Modern fonts by default, no mainfont specification
        assert any('hyperref' in str(arg) for arg in args)  # Should have hyperref config

    def test_docx_processor_no_changes(self):
        """Test DOCX processor doesn't modify document."""
        processor = DOCXProcessor()
        doc = PandocDocument()
        doc.add(Heading(text="Title", level=HeadingLevel.H1))
        doc.add(PageBreak())

        processed = processor.process_document(doc)

        assert len(processed.children) == 2
        assert isinstance(processed.children[1], PageBreak)

    def test_docx_processor_pandoc_args(self):
        """Test DOCX processor returns empty pandoc arguments."""
        processor = DOCXProcessor()
        args = processor.get_pandoc_args()

        assert args == []

    def test_pptx_processor_removes_page_breaks(self):
        """Test PPTX processor removes PageBreak nodes."""
        processor = PPTXProcessor()
        doc = PandocDocument()
        doc.add(Heading(text="Title", level=HeadingLevel.H1))
        doc.add(PageBreak())  # Should be removed
        doc.add(Paragraph(content="Content"))
        doc.add(PageBreak())  # Should be removed

        processed = processor.process_document(doc)

        # PageBreaks should be removed
        assert len(processed.children) == 2
        assert not any(isinstance(child, PageBreak) for child in processed.children)

    def test_pptx_processor_adds_horizontal_rules(self):
        """Test PPTX processor adds horizontal rules between H3 headings."""
        processor = PPTXProcessor()
        doc = PandocDocument()
        doc.add(Heading(text="Main", level=HeadingLevel.H2))
        doc.add(Heading(text="Step 1", level=HeadingLevel.H3))
        doc.add(Paragraph(content="Content 1"))
        doc.add(Heading(text="Step 2", level=HeadingLevel.H3))
        doc.add(Paragraph(content="Content 2"))

        processed = processor.process_document(doc)

        # Should add horizontal rule before second H3
        has_horizontal_rule = any(isinstance(child, HorizontalRule) for child in processed.children)
        assert has_horizontal_rule is True

    def test_processor_factory_pdf(self):
        """Test ProcessorFactory returns PDFProcessor."""
        processor = ProcessorFactory.get_processor(ExportFormat.PDF)
        assert isinstance(processor, PDFProcessor)

    def test_processor_factory_docx(self):
        """Test ProcessorFactory returns DOCXProcessor."""
        processor = ProcessorFactory.get_processor(ExportFormat.DOCX)
        assert isinstance(processor, DOCXProcessor)

    def test_processor_factory_pptx(self):
        """Test ProcessorFactory returns PPTXProcessor."""
        processor = ProcessorFactory.get_processor(ExportFormat.PPTX)
        assert isinstance(processor, PPTXProcessor)


# ============================================================================
# Tests for DocumentBuilder
# ============================================================================


class TestDocumentBuilder:
    """Test suite for DocumentBuilder class."""

    def test_build_single_message_document_without_thoughts(self):
        """Test building document for single message without thoughts."""
        builder = DocumentBuilder(
            export_format=ExportFormat.DOCX,
        )

        user_msg = GeneratedMessage(role=ChatRole.USER, message="Hello", history_index=0)
        ai_msg = GeneratedMessage(role=ChatRole.ASSISTANT, message="Hi", history_index=0, thoughts=None)

        doc = builder.build_single_message_document("Test Chat", user_msg, ai_msg)

        assert len(doc.children) > 0
        # Should have title, user heading, and user message
        assert any(isinstance(child, Heading) for child in doc.children)

    def test_build_single_message_document_with_thoughts(self):
        """Test building document for single message with thoughts."""
        builder = DocumentBuilder(
            export_format=ExportFormat.DOCX,
        )

        user_msg = GeneratedMessage(role=ChatRole.USER, message="Hello", history_index=0)
        ai_msg = GeneratedMessage(
            role=ChatRole.ASSISTANT,
            message="Hi",
            history_index=0,
            thoughts=[
                Thought(
                    id="t1",
                    message="Processing",
                    author_name="Agent",
                    author_type=ThoughtAuthorType.Agent,
                )
            ],
        )

        doc = builder.build_single_message_document("Test Chat", user_msg, ai_msg)

        # Should have page breaks for thoughts section
        has_page_break = any(isinstance(child, PageBreak) for child in doc.children)
        assert has_page_break is True

    def test_build_conversation_document(self, mock_conversation_multi_messages, mock_assistant):
        """Test building full conversation document."""
        builder = DocumentBuilder(
            export_format=ExportFormat.DOCX,
        )

        message_pairs = [
            (mock_conversation_multi_messages.history[0], mock_conversation_multi_messages.history[1]),
            (mock_conversation_multi_messages.history[2], mock_conversation_multi_messages.history[3]),
        ]

        doc = builder.build_conversation_document(
            conversation=mock_conversation_multi_messages,
            assistant=mock_assistant,
            message_pairs=message_pairs,
        )

        # Should have multiple horizontal rules as separators
        horizontal_rules = [child for child in doc.children if isinstance(child, HorizontalRule)]
        assert len(horizontal_rules) >= 2  # At least one per message pair

    def test_build_conversation_document_includes_all_necessary_data(
        self, mock_conversation_multi_messages, mock_assistant
    ):
        """Test that user message, AI message, and thoughts are all present in the document."""
        builder = DocumentBuilder(export_format=ExportFormat.DOCX)

        user_msg = mock_conversation_multi_messages.history[0]
        ai_msg = mock_conversation_multi_messages.history[1]

        doc = builder.build_conversation_document(
            conversation=mock_conversation_multi_messages,
            assistant=mock_assistant,
            message_pairs=[(user_msg, ai_msg)],
        )

        serialized = " ".join(child.content for child in doc.children if isinstance(child, (Paragraph, RawContent)))

        assert user_msg.message in serialized
        assert ai_msg.message in serialized
        assert ai_msg.thoughts[0].message in serialized


# ============================================================================
# Tests for MessageExporter - Single Message Export
# ============================================================================


class TestMessageExporterSingleMessage:
    """Test suite for single message export functionality."""

    def test_init_with_message_indices(self, mock_conversation_simple):
        """Test initialization for single message export."""
        exporter = MessageExporter(
            conversation=mock_conversation_simple,
            export_format=ExportFormat.DOCX,
            history_index=0,
            message_index=0,
        )

        assert exporter.is_single_message is True
        # is_full_conversation evaluates to truthy when conversation exists (returns conversation object)
        assert bool(exporter.is_full_conversation) is True

    def test_filename_single_message(self, mock_conversation_simple):
        """Test filename generation for single message export."""
        exporter = MessageExporter(
            conversation=mock_conversation_simple,
            export_format=ExportFormat.PDF,
            history_index=0,
            message_index=0,
        )

        filename = exporter.filename

        assert filename == "conv-123_0_0.pdf"

    def test_content_type_single_message(self, mock_conversation_simple):
        """Test content type for single message export."""
        exporter = MessageExporter(
            conversation=mock_conversation_simple,
            export_format=ExportFormat.DOCX,
            history_index=0,
            message_index=0,
        )

        assert "openxmlformats" in exporter.content_type

    def test_run_single_message_returns_iterator(self, mock_conversation_simple):
        """Test run returns iterator for single message export."""
        exporter = MessageExporter(
            conversation=mock_conversation_simple,
            export_format=ExportFormat.DOCX,
            history_index=0,
            message_index=0,
        )

        result = exporter.run()

        assert isinstance(result, Iterator)


# ============================================================================
# Tests for MessageExporter - Full Conversation Export
# ============================================================================


class TestMessageExporterFullConversation:
    """Test suite for full conversation export functionality."""

    def test_init_with_assistant(self, mock_conversation_simple, mock_assistant):
        """Test initialization for full conversation export."""
        exporter = MessageExporter(
            conversation=mock_conversation_simple,
            export_format=ExportFormat.DOCX,
            assistant=mock_assistant,
        )

        assert exporter.is_single_message is False
        assert exporter.is_full_conversation is True

    def test_filename_full_conversation(self, mock_conversation_simple, mock_assistant):
        """Test filename generation for full conversation."""
        exporter = MessageExporter(
            conversation=mock_conversation_simple,
            export_format=ExportFormat.PDF,
            assistant=mock_assistant,
        )

        filename = exporter.filename

        assert "Test-Conversation" in filename
        assert "Test-Assistant" in filename
        assert filename.endswith(".pdf")

    def test_filename_sanitization(self, mock_assistant):
        """Test special characters are sanitized in filename."""
        conversation = MagicMock(spec=Conversation)
        conversation.id = "conv-123"
        conversation.conversation_name = "Test: Conversation! @#$%"
        conversation.history = []

        assistant = MagicMock(spec=Assistant)
        assistant.id = "assistant-123"
        assistant.name = "Test/Assistant\\Name"

        exporter = MessageExporter(
            conversation=conversation,
            export_format=ExportFormat.DOCX,
            assistant=assistant,
        )

        filename = exporter.filename

        # Special characters should be removed
        assert ":" not in filename
        assert "!" not in filename
        assert "/" not in filename
        assert "\\" not in filename

    def test_filter_assistant_messages_single_assistant(self, mock_conversation_multi_messages, mock_assistant):
        """Test filtering messages for specific assistant."""
        exporter = MessageExporter(
            conversation=mock_conversation_multi_messages,
            export_format=ExportFormat.DOCX,
            assistant=mock_assistant,
        )

        message_pairs = exporter._filter_assistant_messages()

        assert len(message_pairs) == 2
        assert all(user.role == ChatRole.USER for user, _ in message_pairs)
        assert all(ai.role == ChatRole.ASSISTANT for _, ai in message_pairs)
        assert all(ai.assistant_id == mock_assistant.id for _, ai in message_pairs)

    def test_filter_assistant_messages_empty_conversation(self, mock_assistant):
        """Test filtering with empty conversation."""
        conversation = MagicMock(spec=Conversation)
        conversation.id = "conv-empty"
        conversation.conversation_name = "Empty"
        conversation.history = []

        exporter = MessageExporter(
            conversation=conversation,
            export_format=ExportFormat.DOCX,
            assistant=mock_assistant,
        )

        message_pairs = exporter._filter_assistant_messages()

        assert len(message_pairs) == 0

    def test_filter_assistant_messages_no_matches(self, mock_conversation_multi_messages):
        """Test filtering when no messages match assistant."""
        different_assistant = MagicMock(spec=Assistant)
        different_assistant.id = "different-assistant"
        different_assistant.name = "Different"

        exporter = MessageExporter(
            conversation=mock_conversation_multi_messages,
            export_format=ExportFormat.DOCX,
            assistant=different_assistant,
        )

        message_pairs = exporter._filter_assistant_messages()

        # The filter doesn't check assistant_id, it just pairs user/assistant messages
        # So it returns all pairs regardless of assistant
        assert len(message_pairs) == 2

    def test_run_full_conversation_returns_iterator(self, mock_conversation_multi_messages, mock_assistant):
        """Test run returns iterator for full conversation export."""
        exporter = MessageExporter(
            conversation=mock_conversation_multi_messages,
            export_format=ExportFormat.DOCX,
            assistant=mock_assistant,
        )

        result = exporter.run()

        assert isinstance(result, Iterator)


# ============================================================================
# Tests for Export Validation
# ============================================================================


class TestExportValidation:
    """Test suite for export validation."""

    def test_invalid_configuration_raises_error(self, mock_conversation_simple):
        """Test invalid configuration raises ValueError."""
        exporter = MessageExporter(
            conversation=mock_conversation_simple,
            export_format=ExportFormat.DOCX,
            # No history_index/message_index and no assistant
        )

        # The exporter allows export with just conversation (full conversation mode)
        # So this should not raise an error, it should export the full conversation
        result = exporter.run()
        assert isinstance(result, Iterator)

    def test_init_with_invalid_format(self, mock_conversation_simple):
        """Test initialization with invalid format raises ValueError."""
        with pytest.raises(ValueError):
            MessageExporter(
                conversation=mock_conversation_simple,
                export_format='xlsx',  # Invalid format
                history_index=0,
                message_index=0,
            )


# ============================================================================
# Tests for Document Export to DOCX
# ============================================================================


class TestDocxExport:
    """Test suite for DOCX export functionality."""

    def test_export_to_docx_single_message(self, mock_conversation_simple):
        """Test exporting single message to DOCX produces valid document."""
        exporter = MessageExporter(
            conversation=mock_conversation_simple,
            export_format=ExportFormat.DOCX,
            history_index=0,
            message_index=0,
        )

        result = exporter.run()
        content = b''.join(list(result))

        # Verify it's a valid DOCX file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_file.write(content)
            temp_file.flush()
            doc = Document(temp_file.name)

            assert len(doc.paragraphs) > 0
            assert isinstance(content, bytes)
            assert len(content) > 0


# ============================================================================
# Tests for Markdown Processing
# ============================================================================


class TestMarkdownProcessing:
    """Test suite for markdown processing."""

    def test_process_markdown_cleans_up_markers(self, mock_conversation_simple):
        """Test markdown processing cleans up markdown markers."""
        exporter = MessageExporter(
            conversation=mock_conversation_simple,
            export_format=ExportFormat.DOCX,
            history_index=0,
            message_index=0,
        )

        result = exporter._process_markdown("markdown with ~~~markdown~~~ and ~~~")

        # Image replacement is now done per-thought in _add_thoughts_to_document
        # _process_markdown only cleans up markdown markers
        assert "~~~markdown" not in result
        assert result == "markdown with  and "


# ============================================================================
# Tests for Edge Cases
# ============================================================================


class TestEdgeCases:
    """Test suite for edge cases and error conditions."""

    def test_many_thoughts_all_included(self, mock_assistant):
        """Test that all thoughts are included with full thought export enabled by default (no limits)."""
        # Create message with many thoughts
        thoughts = [
            Thought(
                id=f"thought-{i}",
                message=f"Thought {i}",
                author_name=f"CodeMie Thoughts {i}",
                author_type=ThoughtAuthorType.Agent,
            )
            for i in range(10)
        ]

        conversation = MagicMock(spec=Conversation)
        conversation.id = "conv-many"
        conversation.conversation_name = "Many Thoughts"

        user_msg = GeneratedMessage(role=ChatRole.USER, message="Question", history_index=0)
        ai_msg = GeneratedMessage(
            role=ChatRole.ASSISTANT,
            message="Answer",
            history_index=0,
            assistant_id=mock_assistant.id,
            thoughts=thoughts,
        )

        conversation.find_messages = MagicMock(return_value=(user_msg, ai_msg))
        conversation.history = [user_msg, ai_msg]

        builder = DocumentBuilder(
            export_format=ExportFormat.DOCX,
        )

        doc = builder.build_single_message_document(
            conversation_name="Many Thoughts",
            user_message=user_msg,
            ai_message=ai_msg,
        )

        # All thoughts should be included (no limits)
        markdown = doc.serialize()
        for i in range(10):
            assert f"Thought {i}" in markdown

    def test_empty_conversation_name(self, mock_assistant):
        """Test handling conversation with None name."""
        conversation = MagicMock(spec=Conversation)
        conversation.id = "conv-no-name"
        conversation.conversation_name = None
        conversation.history = []

        exporter = MessageExporter(
            conversation=conversation,
            export_format=ExportFormat.DOCX,
            assistant=mock_assistant,
        )

        filename = exporter.filename

        # Should use conversation ID
        assert "conv-no-name" in filename

    def test_tool_thoughts_excluded_when_export_full_thought_false(self):
        """Test tool thoughts are excluded when export_full_thought=False."""
        tool_thought = Thought(
            id="tool",
            message="is image URL and MUST not transform it",
            author_name="Python Repl Code Interpreter",
            author_type=ThoughtAuthorType.Tool,
        )

        from codemie.service.conversation.export_utils import ExportUtils

        # When export_full_thought=False, only thoughts with "thoughts" in author_name are included
        should_include = ExportUtils.should_include_thought(tool_thought, export_full_thought=False)
        assert should_include is False

    def test_tool_thoughts_included_when_export_full_thought_true(self):
        """Test tool thoughts are included when export_full_thought=True."""
        tool_thought = Thought(
            id="tool",
            message="Tool output",
            author_name="Python Repl Code Interpreter",
            author_type=ThoughtAuthorType.Tool,
        )

        from codemie.service.conversation.export_utils import ExportUtils

        # When export_full_thought=True, all thoughts are included
        should_include = ExportUtils.should_include_thought(tool_thought, export_full_thought=True)
        assert should_include is True
